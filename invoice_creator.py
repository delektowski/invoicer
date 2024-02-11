from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from count_invoice import InvoiceCounter
from number_to_word import number_to_word


class InvoiceCreator(InvoiceCounter):
    doc = Document("invoice_template.docx")
    properties_excluded_align_center = ["invoice_number", "invoice_date", "invoice_pay_date"]

    def __init__(
        self,
        invoice_number,
        invoice_date,
        invoice_pay_date,
        invoice_hour_rate,
        invoice_hours_number,
    ) -> None:
        super().__init__(invoice_hour_rate, invoice_hours_number)

        # Set fields values
        self.invoice_number_data = invoice_number
        self.invoice_date_data = invoice_date
        self.invoice_pay_date_data = invoice_pay_date
        self.invoice_hour_rate_data = invoice_hour_rate
        self.invoice_hours_number_data = invoice_hours_number

        # Set fields references in docx
        self.invoice_number = self.doc.tables[0].cell(0, 7)
        self.invoice_date = self.doc.tables[0].cell(1, 7)
        self.invoice_pay_date = self.doc.tables[0].cell(2, 7)
        self.invoice_hour_rate = self.doc.tables[0].cell(14, 6)
        self.invoice_hours_number = self.doc.tables[0].cell(14, 8)
        self.invoice_netto_value = self.doc.tables[0].cell(15, 8)
        self.invoice_vat_value = self.doc.tables[0].cell(16, 8)
        self.invoice_brutto_value = self.doc.tables[0].cell(17, 8)
        self.invoice_number_word = self.doc.tables[0].cell(19, 8)

    def write_invoice_fields(self):
        self.invoice_number.text = self.invoice_number_data
        self.invoice_date.text = self.invoice_date_data
        self.invoice_pay_date.text = self.invoice_pay_date_data
        self.invoice_hour_rate.text = self.invoice_hour_rate_data
        self.invoice_hours_number.text = self.invoice_hours_number_data
        self.invoice_netto_value.text = self.get_netto_value()
        self.invoice_vat_value.text = self.get_vat_value()
        self.invoice_brutto_value.text = self.get_brutto_value()
        self.invoice_number_word.text = number_to_word(float(self.get_brutto_value()))

        self.handle_table_cells_style()
        self.save_docx()

    def handle_table_cells_style(self):
        for class_property in vars(self).items():
            property_name = str(class_property[0])
            property_value = class_property[1]
            property_value_type = str(type(property_value))
            print(property_name)

            if "Cell" in property_value_type:
                for paragraph in property_value.paragraphs:
                    self.set_paragraph_alignment(property_name, paragraph)
                    self.set_paragraph_style(paragraph)

    def set_paragraph_alignment(self, property_name, paragraph):
        if property_name not in self.properties_excluded_align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_paragraph_style(self, paragraph):
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = "Arial"

    def save_docx(self):
        self.doc.save("invoice_template_modyfied.docx")
